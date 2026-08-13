from __future__ import annotations

import html
import re
from io import BytesIO
from pathlib import Path
from typing import Literal

from docx import Document
from docx.shared import Inches
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as ReportLabImage,
)
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from app.application.use_cases.export_report import ExportableReport

ExportFormat = Literal["markdown", "pdf", "docx"]

BACKEND_ROOT = Path(__file__).resolve().parents[3]


def _safe_asset_path(file_path: str) -> Path | None:
    candidate = Path(file_path)

    if not candidate.is_absolute():
        candidate = BACKEND_ROOT / candidate

    try:
        resolved = candidate.resolve()
        resolved.relative_to(BACKEND_ROOT.resolve())
    except ValueError:
        return None

    if not resolved.is_file():
        return None

    return resolved


def _asset_paths(exportable: ExportableReport) -> list[tuple[object, Path]]:
    resolved_assets = []

    for asset in exportable.assets:
        asset_path = _safe_asset_path(asset.file_path)

        if asset_path is not None:
            resolved_assets.append((asset, asset_path))

    return resolved_assets


def _unique_citations(exportable: ExportableReport) -> list[dict]:
    unique: dict[str, dict] = {}

    for citation in exportable.citations:
        source_id = str(citation["source_id"])

        if source_id not in unique:
            unique[source_id] = citation

    return list(unique.values())


def _normalize_pdf_text(value: str) -> str:
    replacements = str.maketrans(
        {
            "\u2010": "-",
            "\u2011": "-",
            "\u2012": "-",
            "\u2013": "-",
            "\u2014": "-",
            "\u2018": "'",
            "\u2019": "'",
            "\u201c": '"',
            "\u201d": '"',
            "\u2026": "...",
            "\u00a0": " ",
            "\u2022": "-",
            "\u200b": "",
        }
    )

    normalized = value.translate(replacements)

    return normalized.encode("ascii", "ignore").decode("ascii")


def _plain_text(value: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", value)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    text = _normalize_pdf_text(text)

    return html.escape(text)


def _citation_lines(exportable: ExportableReport) -> list[tuple[str, str, str]]:
    lines = []

    for index, citation in enumerate(_unique_citations(exportable), start=1):
        marker = citation.get("inline_marker") or f"[Source {index}]"
        title = citation.get("source_title") or "Untitled source"
        url = citation.get("source_url") or ""

        lines.append((marker, title, url))

    return lines


def render_markdown(exportable: ExportableReport) -> bytes:
    report = exportable.report
    sections = [report.content_markdown.strip()]

    citations = _citation_lines(exportable)

    if citations:
        sections.append("## References\n")

        for marker, title, url in citations:
            sections.append(
                f"{marker} {title}\n\n"
                f"URL: {url}\n",
            )

    assets = _asset_paths(exportable)

    if assets:
        sections.append("## Report Assets\n")

        for asset, asset_path in assets:
            caption = asset.caption or asset.asset_type
            sections.append(
                f"- **{caption}**: `{asset_path.name}`",
            )

    return "\n\n".join(sections).encode("utf-8")


def _append_pdf_asset(
    story: list,
    asset: object,
    asset_path: Path,
) -> None:
    caption = getattr(asset, "caption", None) or getattr(
        asset,
        "asset_type",
        "Report asset",
    )

    story.append(Spacer(1, 0.15 * inch))
    story.append(
        Paragraph(
            _plain_text(caption),
            ParagraphStyle(
                "AssetCaption",
                parent=getSampleStyleSheet()["BodyText"],
                fontName="Helvetica-Bold",
                fontSize=10,
                leading=13,
            ),
        )
    )

    image = ReportLabImage(str(asset_path))
    image.drawWidth = 6.2 * inch
    image.drawHeight = 3.8 * inch
    image.hAlign = "CENTER"

    story.append(image)


def render_pdf(exportable: ExportableReport) -> bytes:
    report = exportable.report

    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title="Deepcite Research Report",
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "ExportTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontSize=18,
        leading=22,
        spaceAfter=16,
    )

    heading_style = ParagraphStyle(
        "ExportHeading",
        parent=styles["Heading2"],
        fontSize=13,
        leading=17,
        spaceBefore=10,
        spaceAfter=6,
    )

    body_style = ParagraphStyle(
        "ExportBody",
        parent=styles["BodyText"],
        fontSize=10.5,
        leading=15.5,
        spaceAfter=9,
    )

    bullet_style = ParagraphStyle(
        "ExportBullet",
        parent=body_style,
        leftIndent=14,
        firstLineIndent=-8,
    )

    story = [
        Paragraph("Deepcite Research Report", title_style),
    ]

    for raw_line in report.content_markdown.splitlines():
        line = raw_line.strip()

        if not line:
            story.append(Spacer(1, 0.08 * inch))
            continue

        if line.startswith("!["):
            continue

        if line.startswith("### "):
            story.append(
                Paragraph(
                    _plain_text(line[4:]),
                    heading_style,
                )
            )
        elif line.startswith("## "):
            story.append(
                Paragraph(
                    _plain_text(line[3:]),
                    heading_style,
                )
            )
        elif line.startswith("# "):
            story.append(
                Paragraph(
                    _plain_text(line[2:]),
                    title_style,
                )
            )
        elif line.startswith("- ") or line.startswith("* "):
            story.append(
                Paragraph(
                    f"- {_plain_text(line[2:])}",
                    bullet_style,
                )
            )
        else:
            story.append(
                Paragraph(
                    _plain_text(line),
                    body_style,
                )
            )

    citations = _citation_lines(exportable)

    if citations:
        story.append(Spacer(1, 0.2 * inch))
        story.append(
            Paragraph("References", heading_style),
        )

        for marker, title, url in citations:
            story.append(
                Paragraph(
                    _plain_text(f"{marker} {title}"),
                    body_style,
                )
            )

            if url:
                story.append(
                    Paragraph(
                        _plain_text(f"URL: {url}"),
                        body_style,
                    )
                )

    assets = _asset_paths(exportable)

    if assets:
        story.append(PageBreak())
        story.append(
            Paragraph("Report Assets", heading_style),
        )

        for asset, asset_path in assets:
            _append_pdf_asset(story, asset, asset_path)

    document.build(story)

    return output.getvalue()


def _add_docx_heading(
    document: Document,
    text: str,
    level: int,
) -> None:
    document.add_heading(text, level=min(level, 3))


def render_docx(exportable: ExportableReport) -> bytes:
    report = exportable.report
    document = Document()

    document.add_heading("Deepcite Research Report", level=0)

    image_pattern = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")

    for raw_line in report.content_markdown.splitlines():
        line = raw_line.strip()

        if not line:
            document.add_paragraph("")
            continue

        image_match = image_pattern.fullmatch(line)

        if image_match:
            image_path = _safe_asset_path(image_match.group(2))

            if image_path is not None:
                document.add_picture(
                    str(image_path),
                    width=Inches(6.0),
                )

            continue

        if line.startswith("### "):
            _add_docx_heading(document, line[4:], 3)
        elif line.startswith("## "):
            _add_docx_heading(document, line[3:], 2)
        elif line.startswith("# "):
            _add_docx_heading(document, line[2:], 1)
        elif line.startswith("- ") or line.startswith("* "):
            document.add_paragraph(
                line[2:],
                style="List Bullet",
            )
        else:
            document.add_paragraph(
                re.sub(
                    r"\[([^\]]+)\]\([^)]+\)",
                    r"\1",
                    line,
                )
            )

    citations = _citation_lines(exportable)

    if citations:
        document.add_heading("References", level=1)

        for marker, title, url in citations:
            document.add_paragraph(
                f"{marker} {title}",
            )

            if url:
                document.add_paragraph(
                    f"URL: {url}",
                )

    assets = _asset_paths(exportable)

    if assets:
        document.add_page_break()
        document.add_heading("Report Assets", level=1)

        for asset, asset_path in assets:
            caption = getattr(asset, "caption", None) or getattr(
                asset,
                "asset_type",
                "Report asset",
            )

            document.add_paragraph(caption)
            document.add_picture(
                str(asset_path),
                width=Inches(6.0),
            )

    output = BytesIO()
    document.save(output)

    return output.getvalue()


def render_report_export(
    exportable: ExportableReport,
    export_format: ExportFormat,
) -> tuple[bytes, str, str]:
    if export_format == "markdown":
        return (
            render_markdown(exportable),
            "text/markdown; charset=utf-8",
            "md",
        )

    if export_format == "pdf":
        return (
            render_pdf(exportable),
            "application/pdf",
            "pdf",
        )

    if export_format == "docx":
        return (
            render_docx(exportable),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )

    raise ValueError(f"Unsupported export format: {export_format}")